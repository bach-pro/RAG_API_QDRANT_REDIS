from __future__ import annotations

from io import BytesIO

from rag_app.data_loader import (
    UploadedFileData,
    chunk_documents,
    clean_html_text,
    load_uploaded_documents,
    load_uploaded_files,
)


def test_chunk_documents_keeps_short_article_whole():
    from rag_app.models import KnowledgeDocument

    document = KnowledgeDocument(id="DOC-7", text="Short document", metadata={"doc_id": "DOC-7"})
    chunks = chunk_documents([document], chunk_size=1000, chunk_overlap=100)

    assert len(chunks) == 1
    assert chunks[0].id == "DOC-7::0"
    assert chunks[0].metadata["chunk_count"] == 1


def test_chunk_documents_splits_long_text():
    from rag_app.models import KnowledgeDocument

    document = KnowledgeDocument(id="DOC-1", text="A" * 2500, metadata={"doc_id": "DOC-1"})
    chunks = chunk_documents([document], chunk_size=1000, chunk_overlap=100)

    assert len(chunks) >= 3
    assert chunks[0].metadata["chunk_count"] == len(chunks)
    assert chunks[0].id == "DOC-1::0"


def test_clean_html_text_removes_article_markup():
    text = clean_html_text(
        "<p><strong>Mô tả:</strong> Lỗi đăng nhập</p>"
        "<p><img src='x.png' />"
        "<strong>Cách xử lý:</strong> Cấp lại license</p>"
    )

    assert "Mô tả: Lỗi đăng nhập" in text
    assert "Cách xử lý: Cấp lại license" in text
    assert "<p>" not in text
    assert "<img" not in text


def test_load_uploaded_documents_normalizes_generic_content_and_metadata():
    docs = load_uploaded_documents(
        [
            {
                "id": "DOC-policy-7",
                "title": "  Chính sách nghỉ phép  ",
                "text": "<p>Nhân viên có 12 ngày phép.</p>",
                "source_type": "policy",
                "metadata": {"department": "HR", "year": 2026},
            }
        ],
        source_label="api_upload",
    )

    assert len(docs) == 1
    assert docs[0].id == "DOC-policy-7"
    assert docs[0].metadata["document_id"] == "policy-7"
    assert docs[0].metadata["department"] == "HR"
    assert docs[0].metadata["year"] == 2026
    assert docs[0].metadata["source"] == "api_upload"
    assert "Title: Chính sách nghỉ phép" in docs[0].text
    assert "Content: Nhân viên có 12 ngày phép." in docs[0].text


def test_load_uploaded_docx_extracts_paragraphs_tables_and_metadata():
    from docx import Document

    docx = Document()
    docx.core_properties.title = "Quy định nội bộ"
    docx.add_paragraph("Nhân viên có 12 ngày phép mỗi năm.")
    table = docx.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Đối tượng"
    table.cell(0, 1).text = "Toàn bộ nhân viên"
    buffer = BytesIO()
    docx.save(buffer)

    docs = load_uploaded_files(
        [
            UploadedFileData(
                filename="policy.docx",
                content=buffer.getvalue(),
                content_type=(
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
            )
        ]
    )

    assert len(docs) == 1
    assert docs[0].metadata["filename"] == "policy.docx"
    assert docs[0].metadata["source_type"] == "docx"
    assert docs[0].metadata["file_size"] == len(buffer.getvalue())
    assert "Title: Quy định nội bộ" in docs[0].text
    assert "Nhân viên có 12 ngày phép mỗi năm." in docs[0].text
    assert "Đối tượng | Toàn bộ nhân viên" in docs[0].text


def test_load_uploaded_pdf_keeps_page_numbers(monkeypatch):
    import pypdf

    class FakePage:
        def __init__(self, text: str) -> None:
            self.text = text

        def extract_text(self) -> str:
            return self.text

    class FakeReader:
        is_encrypted = False
        pages = [FakePage("Trang một"), FakePage("Trang hai")]

    monkeypatch.setattr(pypdf, "PdfReader", lambda *_args, **_kwargs: FakeReader())

    docs = load_uploaded_files(
        [UploadedFileData(filename="guide.pdf", content=b"fake-pdf")]
    )

    assert len(docs) == 2
    assert docs[0].metadata["page_number"] == 1
    assert docs[0].metadata["page_count"] == 2
    assert docs[1].metadata["page_number"] == 2
    assert docs[0].metadata["filename"] == "guide.pdf"
