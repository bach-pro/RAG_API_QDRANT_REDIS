from __future__ import annotations

import logging
from io import BytesIO, StringIO

from fastapi import FastAPI
from fastapi.testclient import TestClient

from rag_api.dependencies.services import (
    get_chat_memory,
    get_ingest_request_store,
    get_job_service,
    get_rag_service,
)
from rag_api.main import create_app
from rag_api.middleware.request_id import RequestIDMiddleware
from rag_api.middleware.request_logging import RequestLoggingMiddleware
from rag_api.services.ingest_processor import IngestJobProcessor
from rag_api.services.ingest_request_store import IngestRequestStore
from rag_api.services.job_service import JobService
from rag_app.config import AppConfig
from rag_app.models import RagResponse, RagStreamResponse, RetrievedDocument
from rag_app.services import BotDeleteResult, IngestResult


class FakeRagService:
    def __init__(self) -> None:
        self.config = AppConfig()
        self.answer_calls = []
        self.answer_stream_calls = []
        self.ingest_documents_calls = []
        self.delete_bot_calls = []

    def count(self) -> int:
        return 7

    def answer(self, **kwargs):
        self.answer_calls.append(kwargs)
        return RagResponse(
            answer="Renew the license file.",
            mode=kwargs["mode"],
            diagnostics={"test": True},
            sources=[
                RetrievedDocument(
                    id="DOC-1::0",
                    text="Document text preview",
                    metadata={
                        "title": "License error",
                        "document_id": "1",
                        "source_type": "faq",
                    },
                    score=0.91,
                )
            ],
        )

    def answer_stream(self, **kwargs):
        self.answer_stream_calls.append(kwargs)
        return RagStreamResponse(
            chunks=iter(["Renew ", "the license file."]),
            mode=kwargs["mode"],
            diagnostics={"test": True},
            sources=[
                RetrievedDocument(
                    id="DOC-1::0",
                    text="Document text preview",
                    metadata={
                        "title": "License error",
                        "document_id": "1",
                        "source_type": "faq",
                    },
                    score=0.91,
                )
            ],
        )

    def ingest_documents(self, **kwargs):
        self.ingest_documents_calls.append(kwargs)
        progress_callback = kwargs.get("progress_callback")
        if progress_callback is not None:
            progress_callback({"event": "documents_loaded", "documents": 1})
            progress_callback({"event": "batch_done", "start": 0, "end": 1, "total": 1})
        return IngestResult(
            documents_loaded=1,
            chunks_indexed=1,
            collection_count=11,
            replaced_existing=kwargs.get("doc_id") is not None,
        )

    def delete_bot(self, bot_id):
        self.delete_bot_calls.append(bot_id)
        return BotDeleteResult(bot_id=bot_id, chunks_deleted=7)


class FakeChatMemory:
    def __init__(self) -> None:
        self.messages = {}

    def get_messages(self, conversation_id):
        return list(self.messages.get(conversation_id, []))

    def append_turn(self, conversation_id, question, answer):
        self.messages.setdefault(conversation_id, []).extend(
            [
                {"role": "user", "content": question},
                {"role": "assistant", "content": answer},
            ]
        )

    def clear(self, conversation_id):
        self.messages.pop(conversation_id, None)


def test_health_returns_collection_count():
    app = create_app()
    service = FakeRagService()
    app.dependency_overrides[get_rag_service] = lambda: service
    client = TestClient(app)

    response = client.get("/health")

    assert response.status_code == 200
    assert response.json()["status"] == "ok"
    assert response.json()["vector_count"] == 7
    assert response.json()["chroma_count"] == 7


def test_chat_endpoint_returns_answer_and_sources():
    app = create_app()
    service = FakeRagService()
    app.dependency_overrides[get_rag_service] = lambda: service
    client = TestClient(app)

    response = client.post(
        "/v1/chat",
        json={
            "bot_id": "bot-support",
            "question": "How to fix license?",
            "mode": "Semantic",
            "top_k": 3,
            "fetch_k": 2,
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["bot_id"] == "bot-support"
    assert body["answer"] == "Renew the license file."
    assert body["mode"] == "Semantic"
    assert body["sources"][0]["id"] == "DOC-1::0"
    assert body["sources"][0]["title"] == "License error"
    assert body["sources"][0]["document_id"] == "1"
    assert body["sources"][0]["source_type"] == "faq"
    assert service.answer_calls[0]["bot_id"] == "bot-support"
    assert service.answer_calls[0]["fetch_k"] == 3


def test_chat_stream_endpoint_returns_sse_events():
    app = create_app()
    service = FakeRagService()
    app.dependency_overrides[get_rag_service] = lambda: service
    client = TestClient(app)

    response = client.post(
        "/v1/chat/stream",
        json={
            "bot_id": "bot-support",
            "question": "How to fix license?",
            "mode": "Semantic",
            "top_k": 3,
            "fetch_k": 2,
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/event-stream")
    body = response.text
    assert "event: metadata" in body
    assert "event: token" in body
    assert '"token": "Renew "' in body
    assert "event: done" in body
    assert '"answer": "Renew the license file."' in body
    assert '"bot_id": "bot-support"' in body
    assert service.answer_stream_calls[0]["bot_id"] == "bot-support"
    assert service.answer_stream_calls[0]["fetch_k"] == 3


def test_chat_history_is_namespaced_by_bot_id():
    app = create_app()
    service = FakeRagService()
    memory = FakeChatMemory()
    app.dependency_overrides[get_rag_service] = lambda: service
    app.dependency_overrides[get_chat_memory] = lambda: memory
    client = TestClient(app)

    for bot_id in ("bot-a", "bot-b", "bot-a"):
        response = client.post(
            "/v1/chat",
            json={
                "bot_id": bot_id,
                "question": f"Question for {bot_id}",
                "conversation_id": "shared-conversation",
            },
        )
        assert response.status_code == 200

    assert service.answer_calls[0]["history"] == []
    assert service.answer_calls[1]["history"] == []
    assert len(service.answer_calls[2]["history"]) == 2
    assert len(memory.messages) == 2


def test_document_ingest_endpoint_queues_request_and_worker_completes_it(tmp_path):
    app = create_app()
    service = FakeRagService()
    job_service = JobService()
    request_store = IngestRequestStore(tmp_path)
    app.dependency_overrides[get_rag_service] = lambda: service
    app.dependency_overrides[get_job_service] = lambda: job_service
    app.dependency_overrides[get_ingest_request_store] = lambda: request_store
    client = TestClient(app)

    response = client.post(
        "/v1/ingest/documents",
        json={
            "bot_id": "bot-policy",
            "doc_id": " policy-v1 ",
            "documents": [
                {
                    "title": "Policy A",
                    "text": "Rule content",
                    "source_type": "policy",
                    "metadata": {"lang": "vi"},
                }
            ],
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["status"] == "queued"
    assert body["bot_id"] == "bot-policy"
    assert body["doc_id"] == "policy-v1"
    assert body["job_id"]
    assert service.ingest_documents_calls == []

    queued_job_id = job_service.dequeue(timeout=1)
    assert queued_job_id == body["job_id"]
    IngestJobProcessor(
        service=service,
        job_service=job_service,
        request_store=request_store,
    ).process(queued_job_id)

    status_response = client.get(f"/v1/ingest/{body['job_id']}")
    assert status_response.status_code == 200
    status = status_response.json()
    assert status["status"] == "completed"
    assert status["bot_id"] == "bot-policy"
    assert status["doc_id"] == "policy-v1"
    assert status["documents_loaded"] == 1
    assert status["chunks_indexed"] == 1
    assert status["collection_count"] == 11
    assert status["replaced_existing"] is True
    assert "reset" not in service.ingest_documents_calls[0]
    assert service.ingest_documents_calls[0]["doc_id"] == "policy-v1"
    assert service.ingest_documents_calls[0]["bot_id"] == "bot-policy"
    assert service.ingest_documents_calls[0]["documents"][0].id == "DOC-policy-v1"
    assert service.ingest_documents_calls[0]["documents"][0].metadata["source"] == "api_upload"


def test_file_ingest_endpoint_saves_file_and_worker_completes_it(tmp_path):
    from docx import Document

    docx = Document()
    docx.add_paragraph("Nội dung chính sách từ file Word.")
    buffer = BytesIO()
    docx.save(buffer)

    app = create_app()
    service = FakeRagService()
    job_service = JobService()
    request_store = IngestRequestStore(tmp_path)
    app.dependency_overrides[get_rag_service] = lambda: service
    app.dependency_overrides[get_job_service] = lambda: job_service
    app.dependency_overrides[get_ingest_request_store] = lambda: request_store
    client = TestClient(app)

    response = client.post(
        "/v1/ingest/files",
        data={"bot_id": "bot-policy", "doc_id": "policy-file"},
        files={
            "files": (
                "policy.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    job_id = response.json()["job_id"]
    assert service.ingest_documents_calls == []
    assert request_store.load(job_id)["files"][0]["filename"] == "policy.docx"

    queued_job_id = job_service.dequeue(timeout=1)
    assert queued_job_id == job_id
    IngestJobProcessor(
        service=service,
        job_service=job_service,
        request_store=request_store,
    ).process(queued_job_id)

    status = client.get(f"/v1/ingest/{job_id}").json()
    assert status["status"] == "completed"
    assert status["bot_id"] == "bot-policy"
    assert status["doc_id"] == "policy-file"
    assert status["documents_loaded"] == 1
    call = service.ingest_documents_calls[0]
    assert "reset" not in call
    assert call["doc_id"] == "policy-file"
    assert call["bot_id"] == "bot-policy"
    assert call["documents"][0].metadata["filename"] == "policy.docx"
    assert call["documents"][0].metadata["source"] == "file_upload"


def test_file_ingest_assigns_one_doc_id_to_each_file(tmp_path):
    from docx import Document

    file_contents = []
    for text in ("Chính sách nghỉ phép.", "Hướng dẫn chấm công."):
        document = Document()
        document.add_paragraph(text)
        buffer = BytesIO()
        document.save(buffer)
        file_contents.append(buffer.getvalue())

    app = create_app()
    service = FakeRagService()
    job_service = JobService()
    request_store = IngestRequestStore(tmp_path)
    app.dependency_overrides[get_rag_service] = lambda: service
    app.dependency_overrides[get_job_service] = lambda: job_service
    app.dependency_overrides[get_ingest_request_store] = lambda: request_store
    client = TestClient(app)

    response = client.post(
        "/v1/ingest/files",
        data={
            "bot_id": "bot-policy",
            "doc_ids": ["leave-policy", "timekeeping-guide"],
        },
        files=[
            (
                "files",
                (
                    "leave.docx",
                    file_contents[0],
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
            (
                "files",
                (
                    "timekeeping.docx",
                    file_contents[1],
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
        ],
    )

    assert response.status_code == 200
    body = response.json()
    assert body["doc_id"] is None
    assert body["doc_ids"] == ["leave-policy", "timekeeping-guide"]
    stored_request = request_store.load(body["job_id"])
    assert [item["doc_id"] for item in stored_request["files"]] == body["doc_ids"]

    queued_job_id = job_service.dequeue(timeout=1)
    assert queued_job_id == body["job_id"]
    IngestJobProcessor(
        service=service,
        job_service=job_service,
        request_store=request_store,
    ).process(queued_job_id)

    status = client.get(f"/v1/ingest/{body['job_id']}").json()
    assert status["status"] == "completed"
    assert status["doc_ids"] == ["leave-policy", "timekeeping-guide"]
    assert status["documents_loaded"] == 2
    assert status["chunks_indexed"] == 2
    assert [call["doc_id"] for call in service.ingest_documents_calls] == [
        "leave-policy",
        "timekeeping-guide",
    ]
    assert [
        call["documents"][0].metadata["filename"]
        for call in service.ingest_documents_calls
    ] == ["leave.docx", "timekeeping.docx"]


def test_file_ingest_auto_generates_ids_when_client_omits_them(tmp_path):
    app = create_app()
    service = FakeRagService()
    job_service = JobService()
    request_store = IngestRequestStore(tmp_path)
    app.dependency_overrides[get_rag_service] = lambda: service
    app.dependency_overrides[get_job_service] = lambda: job_service
    app.dependency_overrides[get_ingest_request_store] = lambda: request_store
    client = TestClient(app)
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    single_response = client.post(
        "/v1/ingest/files",
        data={"bot_id": "bot-policy"},
        files={"files": ("single.docx", b"single", content_type)},
    )

    assert single_response.status_code == 200
    single_body = single_response.json()
    assert single_body["doc_id"].startswith("doc-")
    assert single_body["doc_ids"] is None

    job_service.update(single_body["job_id"], status="completed")
    update_response = client.post(
        "/v1/ingest/files",
        data={"bot_id": "bot-policy", "doc_id": single_body["doc_id"]},
        files={"files": ("single-updated.docx", b"updated", content_type)},
    )
    assert update_response.status_code == 200
    assert update_response.json()["doc_id"] == single_body["doc_id"]

    job_service.update(update_response.json()["job_id"], status="completed")
    multi_response = client.post(
        "/v1/ingest/files",
        data={"bot_id": "bot-policy"},
        files=[
            ("files", ("first.docx", b"first", content_type)),
            ("files", ("second.docx", b"second", content_type)),
        ],
    )

    assert multi_response.status_code == 200
    multi_body = multi_response.json()
    assert multi_body["doc_id"] is None
    assert len(multi_body["doc_ids"]) == 2
    assert all(doc_id.startswith("doc-") for doc_id in multi_body["doc_ids"])
    assert len(set(multi_body["doc_ids"])) == 2


def test_file_ingest_rejects_invalid_doc_id_mapping(tmp_path):
    app = create_app()
    service = FakeRagService()
    app.dependency_overrides[get_rag_service] = lambda: service
    app.dependency_overrides[get_ingest_request_store] = lambda: IngestRequestStore(tmp_path)
    client = TestClient(app)
    upload = {
        "files": (
            "policy.docx",
            b"not-read-because-validation-runs-first",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }

    mismatch = client.post(
        "/v1/ingest/files",
        data={"bot_id": "bot-policy", "doc_ids": ["one", "two"]},
        files=upload,
    )
    duplicate = client.post(
        "/v1/ingest/files",
        data={"bot_id": "bot-policy", "doc_ids": ["same", "same"]},
        files=[("files", upload["files"]), ("files", upload["files"])],
    )
    both = client.post(
        "/v1/ingest/files",
        data={"bot_id": "bot-policy", "doc_id": "group", "doc_ids": ["one"]},
        files=upload,
    )

    assert mismatch.status_code == 422
    assert mismatch.json()["detail"] == (
        "doc_ids must contain exactly one ID for each uploaded file."
    )
    assert duplicate.status_code == 422
    assert duplicate.json()["detail"] == "doc_ids must be unique within one upload request."
    assert both.status_code == 422
    assert both.json()["detail"] == "Use either doc_id or doc_ids, not both."


def test_file_ingest_openapi_marks_files_as_binary_and_exposes_doc_ids():
    client = TestClient(create_app())

    spec = client.get("/openapi.json").json()
    request_schema = spec["paths"]["/v1/ingest/files"]["post"]["requestBody"]["content"][
        "multipart/form-data"
    ]["schema"]
    schema_name = request_schema["$ref"].rsplit("/", 1)[-1]
    properties = spec["components"]["schemas"][schema_name]["properties"]

    assert properties["files"]["items"]["format"] == "binary"
    assert "doc_ids" in properties


def test_file_ingest_endpoint_rejects_unsupported_type():
    app = create_app()
    service = FakeRagService()
    app.dependency_overrides[get_rag_service] = lambda: service
    client = TestClient(app)

    response = client.post(
        "/v1/ingest/files",
        data={"bot_id": "bot-policy"},
        files={"files": ("notes.txt", b"plain text", "text/plain")},
    )

    assert response.status_code == 415
    assert service.ingest_documents_calls == []


def test_ingest_endpoints_require_valid_bot_id():
    app = create_app()
    service = FakeRagService()
    app.dependency_overrides[get_rag_service] = lambda: service
    client = TestClient(app)

    document_response = client.post(
        "/v1/ingest/documents",
        json={"documents": [{"text": "content"}]},
    )
    file_response = client.post(
        "/v1/ingest/files",
        files={"files": ("notes.txt", b"plain text", "text/plain")},
    )
    invalid_response = client.post(
        "/v1/ingest/documents",
        json={
            "bot_id": "invalid bot id",
            "documents": [{"text": "content"}],
        },
    )

    assert document_response.status_code == 422
    assert file_response.status_code == 422
    assert invalid_response.status_code == 422
    assert service.ingest_documents_calls == []


def test_empty_doc_id_generates_stable_document_id_and_reset_is_rejected(tmp_path):
    app = create_app()
    service = FakeRagService()
    job_service = JobService()
    request_store = IngestRequestStore(tmp_path)
    app.dependency_overrides[get_rag_service] = lambda: service
    app.dependency_overrides[get_job_service] = lambda: job_service
    app.dependency_overrides[get_ingest_request_store] = lambda: request_store
    client = TestClient(app)

    response = client.post(
        "/v1/ingest/documents",
        json={
            "bot_id": "bot-policy",
            "doc_id": "   ",
            "documents": [{"text": "content"}],
        },
    )
    first_job_id = job_service.dequeue(timeout=1)
    IngestJobProcessor(
        service=service,
        job_service=job_service,
        request_store=request_store,
    ).process(first_job_id)
    second_response = client.post(
        "/v1/ingest/documents",
        json={
            "bot_id": "bot-policy",
            "documents": [{"text": "other content"}],
        },
    )
    second_job_id = job_service.dequeue(timeout=1)
    IngestJobProcessor(
        service=service,
        job_service=job_service,
        request_store=request_store,
    ).process(second_job_id)
    generated_doc_id = response.json()["doc_id"]
    update_response = client.post(
        "/v1/ingest/documents",
        json={
            "bot_id": "bot-policy",
            "doc_id": generated_doc_id,
            "documents": [{"text": "updated content"}],
        },
    )
    update_job_id = job_service.dequeue(timeout=1)
    IngestJobProcessor(
        service=service,
        job_service=job_service,
        request_store=request_store,
    ).process(update_job_id)
    reset_response = client.post(
        "/v1/ingest/documents",
        json={
            "bot_id": "bot-policy",
            "reset": True,
            "documents": [{"text": "content"}],
        },
    )

    assert response.status_code == 200
    assert second_response.status_code == 200
    first_doc_id = response.json()["doc_id"]
    second_doc_id = second_response.json()["doc_id"]
    assert first_doc_id.startswith("doc-")
    assert second_doc_id.startswith("doc-")
    assert first_doc_id != second_doc_id
    assert service.ingest_documents_calls[0]["doc_id"] == first_doc_id
    assert service.ingest_documents_calls[1]["doc_id"] == second_doc_id
    assert update_response.json()["doc_id"] == first_doc_id
    assert service.ingest_documents_calls[2]["doc_id"] == first_doc_id
    first_internal_id = service.ingest_documents_calls[0]["documents"][0].id
    second_internal_id = service.ingest_documents_calls[1]["documents"][0].id
    assert first_internal_id == f"DOC-{first_doc_id.removeprefix('doc-')}"
    assert second_internal_id == f"DOC-{second_doc_id.removeprefix('doc-')}"
    assert first_internal_id != second_internal_id
    assert reset_response.status_code == 422


def test_document_item_id_is_no_longer_accepted():
    app = create_app()
    service = FakeRagService()
    app.dependency_overrides[get_rag_service] = lambda: service
    client = TestClient(app)

    response = client.post(
        "/v1/ingest/documents",
        json={
            "bot_id": "bot-policy",
            "doc_id": "policy-v1",
            "documents": [{"id": "legacy-id", "text": "content"}],
        },
    )
    document_schema = client.get("/openapi.json").json()["components"]["schemas"][
        "IngestDocumentInput"
    ]

    assert response.status_code == 422
    assert "id" not in document_schema["properties"]
    assert service.ingest_documents_calls == []


def test_delete_bot_endpoint_deletes_only_the_requested_bot_data():
    app = create_app()
    service = FakeRagService()
    app.dependency_overrides[get_rag_service] = lambda: service
    client = TestClient(app)

    response = client.delete("/v1/bots/bot-policy")

    assert response.status_code == 200
    assert response.json() == {"bot_id": "bot-policy", "chunks_deleted": 7}
    assert service.delete_bot_calls == ["bot-policy"]


def test_ingest_endpoint_rejects_second_active_job(tmp_path):
    app = create_app()
    service = FakeRagService()
    job_service = JobService()
    active_job = job_service.create_ingest_job("bot-active")
    app.dependency_overrides[get_rag_service] = lambda: service
    app.dependency_overrides[get_job_service] = lambda: job_service
    app.dependency_overrides[get_ingest_request_store] = lambda: IngestRequestStore(tmp_path)
    client = TestClient(app)

    response = client.post(
        "/v1/ingest/documents",
        json={
            "bot_id": "bot-policy",
            "documents": [{"text": "content"}],
        },
    )

    assert response.status_code == 409
    detail = response.json()["detail"]
    assert detail["message"] == "An ingest job is already active."
    assert detail["job_id"] == active_job.job_id
    assert detail["bot_id"] == "bot-active"
    assert detail["status"] == "queued"
    assert service.ingest_documents_calls == []


def test_request_id_and_logging_middleware():
    app = FastAPI()
    app.add_middleware(RequestLoggingMiddleware, logger_name="rag_api.middleware_test")
    app.add_middleware(RequestIDMiddleware)

    @app.get("/ping")
    def ping():
        return {"ok": True}

    client = TestClient(app)
    log_stream = StringIO()
    handler = logging.StreamHandler(log_stream)
    logger = logging.getLogger("rag_api.middleware_test")
    logger.setLevel(logging.INFO)
    logger.addHandler(handler)

    try:
        response = client.get("/ping", headers={"X-Request-ID": "req-123"})
    finally:
        logger.removeHandler(handler)

    assert response.status_code == 200
    assert response.headers["X-Request-ID"] == "req-123"
    assert "GET /ping 200" in log_stream.getvalue()
    assert "request_id=req-123" in log_stream.getvalue()
