from __future__ import annotations

import html
import hashlib
import json
import re
from dataclasses import dataclass
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import Any

from .models import KnowledgeDocument


SUPPORTED_FILE_EXTENSIONS = {".pdf", ".docx"}


@dataclass(frozen=True)
class UploadedFileData:
    filename: str
    content: bytes
    content_type: str | None = None


class _HTMLTextExtractor(HTMLParser):
    block_tags = {
        "address",
        "article",
        "aside",
        "blockquote",
        "br",
        "dd",
        "div",
        "dl",
        "dt",
        "figcaption",
        "figure",
        "footer",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "header",
        "hr",
        "li",
        "main",
        "nav",
        "ol",
        "p",
        "pre",
        "section",
        "table",
        "td",
        "th",
        "tr",
        "ul",
    }

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() == "li":
            self._append_break()
            self._parts.append("- ")
        elif tag.lower() in self.block_tags:
            self._append_break()

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in self.block_tags:
            self._append_break()

    def handle_data(self, data: str) -> None:
        if data:
            self._parts.append(data)

    def _append_break(self) -> None:
        if self._parts and self._parts[-1] != "\n":
            self._parts.append("\n")

    def text(self) -> str:
        return "".join(self._parts)


def clean_vietnamese_text(text: str) -> str:
    """Normalize whitespace while preserving Vietnamese accents and field layout."""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def clean_html_text(text: str) -> str:
    """Convert simple article HTML to plain text suitable for retrieval."""

    text = re.sub(r"<img[^>]*>", "", str(text), flags=re.IGNORECASE)
    if "<" not in text or ">" not in text:
        return clean_vietnamese_text(html.unescape(text))

    parser = _HTMLTextExtractor()
    parser.feed(text)
    parser.close()
    extracted = parser.text()
    extracted = re.sub(r"[ \t]*\n[ \t]*", "\n", extracted)
    extracted = re.sub(r"\n{3,}", "\n\n", extracted)
    return clean_vietnamese_text(extracted)


def _scalar_metadata(value: Any) -> str | int | float | bool | None:
    if value is None:
        return None
    if isinstance(value, (str, int, float, bool)):
        return value
    return json.dumps(value, ensure_ascii=False)


def sanitize_metadata(metadata: dict[str, Any]) -> dict[str, str | int | float | bool]:
    clean: dict[str, str | int | float | bool] = {}
    for key, value in metadata.items():
        scalar = _scalar_metadata(value)
        if scalar is not None:
            clean[str(key)] = scalar
    return clean


def _first_non_empty_value(record: dict[str, Any], keys: tuple[str, ...]) -> str:
    for key in keys:
        value = record.get(key)
        if value not in (None, ""):
            text = clean_vietnamese_text(str(value))
            if text:
                return text
    return ""


def clean_generic_document_record(record: dict[str, Any]) -> dict[str, Any]:
    """Normalize one generic data row for upload-based ingestion."""

    return {
        "id": record.get("id") or record.get("doc_id") or record.get("document_id") or record.get("source_id"),
        "title": _first_non_empty_value(record, ("title", "name", "subject", "instruction", "question")),
        "keywords": _first_non_empty_value(record, ("keywords", "tags", "input", "labels")),
        "content": clean_html_text(
            str(
                record.get("text")
                or record.get("content")
                or record.get("body")
                or record.get("output")
                or record.get("description")
                or ""
            )
        ),
        "source_type": _first_non_empty_value(record, ("source_type", "type", "category", "source")),
        "metadata": dict(record.get("metadata") or {}),
    }


def build_page_content_from_generic_record(record: dict[str, Any]) -> str:
    clean_record = clean_generic_document_record(record)
    parts: list[str] = []
    document_id = clean_record.get("id")
    title = clean_record.get("title")
    keywords = clean_record.get("keywords")
    source_type = clean_record.get("source_type")
    content = clean_record.get("content")

    if document_id not in (None, ""):
        parts.append(f"Document ID: {document_id}")
    if title not in (None, ""):
        parts.append(f"Title: {title}")
    if source_type not in (None, ""):
        parts.append(f"Source Type: {source_type}")
    if keywords not in (None, ""):
        parts.append(f"Keywords: {keywords}")
    if content:
        parts.append(f"Content: {content}")
    return clean_vietnamese_text("\n".join(parts))


def generic_record_to_document(
    record: dict[str, Any],
    source_label: str,
    index: int,
) -> KnowledgeDocument | None:
    text = build_page_content_from_generic_record(record)
    if not text:
        return None

    clean_record = clean_generic_document_record(record)
    document_id = clean_record.get("id")
    if document_id in (None, ""):
        document_id = str(index)
    else:
        document_id = str(document_id).strip()
    if document_id.upper().startswith("DOC-"):
        document_id = document_id[4:]
    doc_id = f"DOC-{document_id}"
    metadata = dict(clean_record.get("metadata") or {})
    metadata.update({
        "doc_id": doc_id,
        "document_id": document_id,
        "title": clean_record.get("title") or "",
        "keywords": clean_record.get("keywords") or "",
        "source": source_label,
        "source_type": clean_record.get("source_type") or "",
        "row_index": index,
    })
    return KnowledgeDocument(
        id=doc_id,
        text=text,
        metadata=sanitize_metadata(metadata),
    )


def load_uploaded_documents(
    documents: list[dict[str, Any]],
    *,
    source_label: str = "uploaded_api",
) -> list[KnowledgeDocument]:
    loaded: list[KnowledgeDocument] = []
    for index, record in enumerate(documents):
        document = generic_record_to_document(record, source_label=source_label, index=index)
        if document is not None:
            loaded.append(document)
    return loaded


def _safe_filename(filename: str) -> str:
    safe_name = Path((filename or "").replace("\x00", "")).name.strip()
    if not safe_name:
        raise ValueError("Uploaded file must have a filename.")
    return safe_name


def _file_identity(filename: str, content: bytes) -> tuple[str, str]:
    content_sha256 = hashlib.sha256(content).hexdigest()
    identity = hashlib.sha256(
        filename.casefold().encode("utf-8") + b"\x00" + content
    ).hexdigest()[:16]
    return f"file-{identity}", content_sha256


def _file_metadata(
    file: UploadedFileData,
    *,
    filename: str,
    content_sha256: str,
) -> dict[str, Any]:
    return {
        "filename": filename,
        "content_type": file.content_type or "application/octet-stream",
        "content_sha256": content_sha256,
        "file_size": len(file.content),
    }


def _load_pdf_file(
    file: UploadedFileData,
    *,
    filename: str,
    source_label: str,
    max_extracted_chars: int,
) -> list[KnowledgeDocument]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(file.content), strict=False)
        if reader.is_encrypted:
            raise ValueError(f"PDF '{filename}' is encrypted and cannot be indexed.")
        page_count = len(reader.pages)
        file_id, content_sha256 = _file_identity(filename, file.content)
        base_metadata = _file_metadata(
            file,
            filename=filename,
            content_sha256=content_sha256,
        )
        documents: list[KnowledgeDocument] = []
        extracted_chars = 0
        for page_index, page in enumerate(reader.pages):
            page_text = clean_vietnamese_text(page.extract_text() or "")
            if not page_text:
                continue
            extracted_chars += len(page_text)
            if extracted_chars > max_extracted_chars:
                raise ValueError(
                    f"Extracted text from '{filename}' exceeds {max_extracted_chars} characters."
                )
            page_number = page_index + 1
            document = generic_record_to_document(
                {
                    "id": f"{file_id}-page-{page_number}",
                    "title": filename,
                    "text": page_text,
                    "source_type": "pdf",
                    "metadata": {
                        **base_metadata,
                        "page_number": page_number,
                        "page_count": page_count,
                    },
                },
                source_label=source_label,
                index=page_index,
            )
            if document is not None:
                documents.append(document)
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Could not read PDF '{filename}': {exc}") from exc

    if not documents:
        raise ValueError(
            f"PDF '{filename}' contains no extractable text. Scanned PDFs require OCR first."
        )
    return documents


def _load_docx_file(
    file: UploadedFileData,
    *,
    filename: str,
    source_label: str,
    max_extracted_chars: int,
) -> list[KnowledgeDocument]:
    from docx import Document

    try:
        docx = Document(BytesIO(file.content))
        blocks = [
            clean_vietnamese_text(paragraph.text)
            for paragraph in docx.paragraphs
            if clean_vietnamese_text(paragraph.text)
        ]
        for table in docx.tables:
            for row in table.rows:
                cells = [clean_vietnamese_text(cell.text) for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    blocks.append(row_text)
        text = clean_vietnamese_text("\n\n".join(blocks))
        if not text:
            raise ValueError(f"DOCX '{filename}' contains no extractable text.")
        if len(text) > max_extracted_chars:
            raise ValueError(
                f"Extracted text from '{filename}' exceeds {max_extracted_chars} characters."
            )

        file_id, content_sha256 = _file_identity(filename, file.content)
        title = clean_vietnamese_text(docx.core_properties.title or "") or filename
        document = generic_record_to_document(
            {
                "id": file_id,
                "title": title,
                "text": text,
                "source_type": "docx",
                "metadata": _file_metadata(
                    file,
                    filename=filename,
                    content_sha256=content_sha256,
                ),
            },
            source_label=source_label,
            index=0,
        )
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Could not read DOCX '{filename}': {exc}") from exc

    if document is None:
        raise ValueError(f"DOCX '{filename}' contains no extractable text.")
    return [document]


def load_uploaded_files(
    files: list[UploadedFileData],
    *,
    source_label: str = "file_upload",
    max_extracted_chars: int = 2_000_000,
) -> list[KnowledgeDocument]:
    if not files:
        raise ValueError("At least one PDF or DOCX file is required.")
    if max_extracted_chars <= 0:
        raise ValueError("max_extracted_chars must be positive.")

    loaded: list[KnowledgeDocument] = []
    for file in files:
        filename = _safe_filename(file.filename)
        extension = Path(filename).suffix.casefold()
        if extension not in SUPPORTED_FILE_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_FILE_EXTENSIONS))
            raise ValueError(f"Unsupported file '{filename}'. Supported types: {supported}.")
        if not file.content:
            raise ValueError(f"Uploaded file '{filename}' is empty.")
        if extension == ".pdf":
            loaded.extend(
                _load_pdf_file(
                    file,
                    filename=filename,
                    source_label=source_label,
                    max_extracted_chars=max_extracted_chars,
                )
            )
        else:
            loaded.extend(
                _load_docx_file(
                    file,
                    filename=filename,
                    source_label=source_label,
                    max_extracted_chars=max_extracted_chars,
                )
            )
    return loaded


def _find_chunk_end(text: str, start: int, max_end: int, min_end: int) -> int:
    if max_end >= len(text):
        return len(text)
    candidates = [
        text.rfind("\n\n", start, max_end),
        text.rfind("\n", start, max_end),
        text.rfind(". ", start, max_end),
        text.rfind("; ", start, max_end),
    ]
    best = max(candidates)
    if best >= min_end:
        return best + 1
    return max_end


def chunk_documents(
    documents: list[KnowledgeDocument],
    chunk_size: int = 1024,
    chunk_overlap: int = 128,
) -> list[KnowledgeDocument]:
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive.")
    if chunk_overlap < 0 or chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be >= 0 and smaller than chunk_size.")

    chunks: list[KnowledgeDocument] = []
    for document in documents:
        text = document.text
        if len(text) <= chunk_size:
            metadata = dict(document.metadata)
            metadata.update({"source_doc_id": document.id, "chunk_index": 0, "chunk_count": 1})
            chunks.append(KnowledgeDocument(id=f"{document.id}::0", text=text, metadata=metadata))
            continue

        doc_chunks: list[KnowledgeDocument] = []
        start = 0
        chunk_index = 0
        while start < len(text):
            min_end = start + max(1, chunk_size // 2)
            end = _find_chunk_end(text, start, min(start + chunk_size, len(text)), min_end)
            chunk_text = clean_vietnamese_text(text[start:end])
            if chunk_text:
                metadata = dict(document.metadata)
                metadata.update(
                    {
                        "source_doc_id": document.id,
                        "chunk_index": chunk_index,
                    }
                )
                doc_chunks.append(
                    KnowledgeDocument(
                        id=f"{document.id}::{chunk_index}",
                        text=chunk_text,
                        metadata=metadata,
                    )
                )
                chunk_index += 1
            if end >= len(text):
                break
            start = max(end - chunk_overlap, start + 1)

        for chunk in doc_chunks:
            chunk.metadata["chunk_count"] = len(doc_chunks)
        chunks.extend(doc_chunks)

    return chunks
